import argparse
from PIL import Image
import pytesseract
from docx import Document
#from fpdf import FPDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.utils import simpleSplit
import os

def extract_text_from_image(image_path):
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img)
    lines = text.split('\n')
    processed_text = ""
    for line in lines:
        if line.strip():  
            processed_text += line.strip() + " "
        else:  
            processed_text = processed_text.strip() + "\n"
    return processed_text.strip()

def save_text_to_docx(text, output_path):
    doc = Document()
    lines = text.split('\n')
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')
    doc.save(output_path)

def save_text_to_pdf(text, output_path):
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    pdfmetrics.registerFont(TTFont('Cambria', 'cambria.ttf'))
    font_size = 12
    c.setFont('Cambria', font_size)

    # Define margins and initial drawing position
    left_margin = 40
    top_margin = height - 40
    line_height = font_size + 4  # Space between lines

    # Split the text into lines that fit within the page width
    max_width = width - 2 * left_margin
    lines = text.split('\n')

    y = top_margin

    for line in lines:
        wrapped_lines = simpleSplit(line, 'Cambria', font_size, max_width)
        for wrapped_line in wrapped_lines:
            if y < 40:  
                c.showPage()
                c.setFont('Cambria', font_size)
                y = top_margin
            c.drawString(left_margin, y, wrapped_line)
            y -= line_height  

        y -= line_height

    c.save()
    """
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    pdfmetrics.registerFont(TTFont('Cambria', 'cambria.ttf'))
    c.setFont('Cambria', 12)

    lines = text.split('\n')
    y = height - 40  
    for line in lines:
        c.drawString(40, y, line)
        y -= 14  

    c.save()
    """
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("Cambria", "", "cambria.ttf", uni=True)
    pdf.set_font("Cambria", size=12)
    lines = text.split('\n')
    for line in lines:
        pdf.multi_cell(0, 10, line)
    pdf.output(output_path)
    """

def convert_image_to_text_doc(image_path, generate_pdf=False):
    text = extract_text_from_image(image_path)
    
    base_filename = os.path.splitext(image_path)[0]
    docx_path = f"{base_filename}.docx"
    pdf_path = f"{base_filename}.pdf"

    save_text_to_docx(text, docx_path)
    
    if generate_pdf:
        save_text_to_pdf(text, pdf_path)

    if generate_pdf:
        return docx_path, pdf_path
    else:
        return docx_path

def main():
    parser = argparse.ArgumentParser(description="Convert image text to docx and optionally a PDF.")
    parser.add_argument("image_path", help="Path to the screenshot image containing text.")
    parser.add_argument("--generate_pdf", action="store_true", help="Generate a PDF copy of the text document.")
    
    args = parser.parse_args()
    
    output_files = convert_image_to_text_doc(args.image_path, args.generate_pdf)
    print("Generated files:", output_files)

if __name__ == "__main__":
    main()
